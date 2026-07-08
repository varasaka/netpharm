"""Agent 11 — Report generator.

Assembles a publication-style report from every upstream table. The narrative
sections (Methods, Results, figure legends, GO/KEGG interpretation, hub-gene
discussion) are drafted by the Claude API from the actual numbers, then rendered
to Markdown, DOCX, HTML, and PDF.

If ANTHROPIC_API_KEY is unset, the agent still produces a complete report using
deterministic templated prose, so the pipeline never hard-fails on the LLM step.

The output_table is a manifest of report files written.
"""
from __future__ import annotations

import json
from pathlib import Path

import pandas as pd

from ..config import Config
from ..db import Store
from .base import BaseAgent


class ReportAgent(BaseAgent):
    name = "report"
    output_table = "report_manifest"

    def run(self, store: Store, config: Config) -> pd.DataFrame:
        cfg = config.section("report")
        facts = self._gather_facts(store, config)
        sections = self._draft_sections(facts, config)
        report_dir = Path(config.output_dir) / "report"
        report_dir.mkdir(parents=True, exist_ok=True)

        md_path = report_dir / "report.md"
        md_path.write_text(self._to_markdown(cfg["title"], sections, facts), encoding="utf-8")
        manifest = [{"format": "markdown", "path": str(md_path)}]

        formats = set(cfg.get("formats", []))
        if "html" in formats:
            manifest.append(self._write_html(report_dir, cfg["title"], sections, facts))
        if "docx" in formats:
            manifest.append(self._write_docx(report_dir, cfg["title"], sections, facts))
        if "pdf" in formats:
            manifest.append(self._write_pdf(report_dir, cfg["title"], sections, facts))

        self.log.info("report written: %s", ", ".join(m["format"] for m in manifest))
        return pd.DataFrame(manifest)

    # ------------------------------------------------------ gather facts
    def _gather_facts(self, store: Store, config: Config) -> dict:
        def maybe(name):
            return store.load_table(name) if store.has_table(name) else pd.DataFrame()

        inter = maybe("intersection_targets")
        venn = maybe("intersection_venn")
        enrich = maybe("enrichment_results")
        hubs = maybe("hub_genes")
        ppi_stats = maybe("ppi_stats")

        top_hubs = hubs.loc[hubs.get("is_hub", False), "gene"].tolist() if not hubs.empty else []
        top_kegg = (
            enrich[enrich["category"] == "KEGG"].sort_values("adj_p_value")["term"].head(10).tolist()
            if not enrich.empty else []
        )
        top_go = (
            enrich[enrich["category"] == "GO:BP"].sort_values("adj_p_value")["term"].head(10).tolist()
            if not enrich.empty else []
        )
        return {
            "plant": config["run"]["plant"],
            "disease": config["run"]["disease"],
            "n_shared_targets": len(inter),
            "venn": venn.to_dict("records")[0] if not venn.empty else {},
            "ppi_stats": ppi_stats.to_dict("records")[0] if not ppi_stats.empty else {},
            "top_hubs": top_hubs,
            "top_kegg": top_kegg,
            "top_go_bp": top_go,
            "n_enriched": len(enrich),
        }

    # ------------------------------------------------------ narrative
    def _draft_sections(self, facts: dict, config: Config) -> dict[str, str]:
        key = config.env("ANTHROPIC_API_KEY")
        if not key:
            self.log.warning("ANTHROPIC_API_KEY unset — using templated narrative.")
            return self._template_sections(facts)
        try:
            import anthropic  # lazy import
            client = anthropic.Anthropic(api_key=key)
            llm = config.section("llm")
            prompt = (
                "You are drafting a network-pharmacology manuscript. Using ONLY the "
                "JSON facts below, write four concise sections as JSON with keys "
                "'methods', 'results', 'enrichment_interpretation', 'hub_discussion'. "
                "Be factual, do not invent numbers, use formal scientific tone.\n\n"
                f"FACTS:\n{json.dumps(facts, indent=2)}"
            )
            msg = client.messages.create(
                model=llm.get("model", "claude-opus-4-8"),
                max_tokens=llm.get("max_tokens", 4000),
                messages=[{"role": "user", "content": prompt}],
            )
            text = "".join(b.text for b in msg.content if b.type == "text")
            text = text.strip().removeprefix("```json").removeprefix("```").removesuffix("```")
            data = json.loads(text)
            base = self._template_sections(facts)
            base.update({k: v for k, v in data.items() if isinstance(v, str)})
            return base
        except Exception as exc:  # noqa: BLE001
            self.log.warning("LLM drafting failed (%s) — templated narrative.", exc)
            return self._template_sections(facts)

    def _template_sections(self, f: dict) -> dict[str, str]:
        return {
            "methods": (
                f"Phytochemicals of {f['plant']} were retrieved and standardized, "
                "screened for drug-likeness (Lipinski, GI absorption, bioavailability, "
                "PAINS), and mapped to protein targets. Disease-associated genes for "
                f"{f['disease']} were collected and intersected with the compound "
                "targets. The shared targets were analyzed by STRING PPI, Enrichr GO/"
                "KEGG/Reactome enrichment, cytoHubba-family hub ranking, and network "
                "visualization in Cytoscape."
            ),
            "results": (
                f"The analysis identified {f['n_shared_targets']} shared targets between "
                f"{f['plant']} constituents and {f['disease']}. The PPI network comprised "
                f"{f.get('ppi_stats', {}).get('nodes', 'n/a')} nodes and "
                f"{f.get('ppi_stats', {}).get('edges', 'n/a')} edges. "
                f"{f['n_enriched']} enriched terms passed the adjusted-p cutoff."
            ),
            "enrichment_interpretation": (
                "Top enriched KEGG pathways: " + (", ".join(f["top_kegg"]) or "n/a") + ". "
                "Top GO biological processes: " + (", ".join(f["top_go_bp"]) or "n/a") + "."
            ),
            "hub_discussion": (
                "Hub genes ranked by MCC: " + (", ".join(f["top_hubs"]) or "n/a") +
                ". These occupy central positions in the PPI network and represent "
                "candidate mechanistic targets."
            ),
            "figure_legends": (
                "Figure 1. Compound–Target network. Figure 2. Compound–Target–Disease "
                "network. Figure 3. PPI network with hub genes highlighted. "
                "Figure 4. Target–Pathway network."
            ),
        }

    # ------------------------------------------------------ renderers
    def _to_markdown(self, title: str, s: dict, f: dict) -> str:
        order = [
            ("Methods", "methods"), ("Results", "results"),
            ("GO / KEGG Interpretation", "enrichment_interpretation"),
            ("Hub Gene Discussion", "hub_discussion"),
            ("Figure Legends", "figure_legends"),
        ]
        parts = [f"# {title}\n", f"**Plant:** {f['plant']}  \n**Disease:** {f['disease']}\n"]
        for heading, key in order:
            parts.append(f"## {heading}\n\n{s.get(key, '')}\n")
        return "\n".join(parts)

    def _write_html(self, d: Path, title: str, s: dict, f: dict) -> dict:
        import markdown  # from the 'markdown' package
        html = markdown.markdown(self._to_markdown(title, s, f))
        path = d / "report.html"
        path.write_text(f"<html><body>{html}</body></html>", encoding="utf-8")
        return {"format": "html", "path": str(path)}

    def _write_docx(self, d: Path, title: str, s: dict, f: dict) -> dict:
        from docx import Document
        doc = Document()
        doc.add_heading(title, level=0)
        doc.add_paragraph(f"Plant: {f['plant']}")
        doc.add_paragraph(f"Disease: {f['disease']}")
        for heading, key in [
            ("Methods", "methods"), ("Results", "results"),
            ("GO / KEGG Interpretation", "enrichment_interpretation"),
            ("Hub Gene Discussion", "hub_discussion"),
            ("Figure Legends", "figure_legends"),
        ]:
            doc.add_heading(heading, level=1)
            doc.add_paragraph(s.get(key, ""))
        path = d / "report.docx"
        doc.save(str(path))
        return {"format": "docx", "path": str(path)}

    def _write_pdf(self, d: Path, title: str, s: dict, f: dict) -> dict:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        path = d / "report.pdf"
        doc = SimpleDocTemplate(str(path), pagesize=A4)
        styles = getSampleStyleSheet()
        flow = [Paragraph(title, styles["Title"]),
                Paragraph(f"Plant: {f['plant']} — Disease: {f['disease']}", styles["Normal"]),
                Spacer(1, 12)]
        for heading, key in [
            ("Methods", "methods"), ("Results", "results"),
            ("GO / KEGG Interpretation", "enrichment_interpretation"),
            ("Hub Gene Discussion", "hub_discussion"),
            ("Figure Legends", "figure_legends"),
        ]:
            flow += [Paragraph(heading, styles["Heading1"]),
                     Paragraph(s.get(key, ""), styles["Normal"]), Spacer(1, 8)]
        doc.build(flow)
        return {"format": "pdf", "path": str(path)}
